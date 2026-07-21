import os
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def convert_md_to_docx(md_file_path: Path, output_paths: list[Path]):
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling defaults
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    content = md_file_path.read_text(encoding='utf-8')
    lines = content.splitlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal in_table, table_lines
        if not table_lines:
            return
        
        # Parse markdown table
        header = [c.strip() for c in table_lines[0].strip('|').split('|')]
        rows = []
        for line in table_lines[2:]:
            if '|' in line:
                rows.append([c.strip() for c in line.strip('|').split('|')])
        
        if header and rows:
            table = doc.add_table(rows=len(rows) + 1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Format header
            for col_idx, text in enumerate(header):
                cell = table.cell(0, col_idx)
                cell.text = text
                set_cell_background(cell, "1F4E78")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
            # Format data rows
            for row_idx, row_data in enumerate(rows):
                fill_color = "F2F4F7" if row_idx % 2 == 1 else "FFFFFF"
                for col_idx, text in enumerate(row_data[:len(header)]):
                    cell = table.cell(row_idx + 1, col_idx)
                    cell.text = text
                    set_cell_background(cell, fill_color)

            doc.add_paragraph()  # Spacing
        table_lines = []
        in_table = False

    def add_formatted_paragraph(text, style=None, space_after=6, space_before=0, is_quote=False):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if is_quote:
            p.paragraph_format.left_indent = Inches(0.4)

        # Parse inline formatting (**bold**, `code`, links)
        tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                run = p.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('`') and token.endswith('`'):
                run = p.add_run(token[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x88, 0x11, 0x11)
            elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
                match = re.match(r'\[(.*?)\]\((.*?)\)', token)
                if match:
                    link_text = match.group(1)
                    run = p.add_run(link_text)
                    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
                    run.underline = True
            else:
                p.add_run(token)

    for line in lines:
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_block_lines)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F4F6F8")
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                doc.add_paragraph()
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Markdown Table handling
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            flush_table()

        # Headers
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_formatted_paragraph(stripped[2:], style='List Bullet', space_after=3)
        elif stripped.startswith('> '):
            add_formatted_paragraph(stripped[2:], space_after=6, is_quote=True)
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run('━' * 55)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif stripped:
            add_formatted_paragraph(stripped, space_after=6)

    if in_table:
        flush_table()

    for out_path in output_paths:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        print(f"Saved: {out_path}")

if __name__ == '__main__':
    src = Path(r"c:\Users\secha\OpsControl\docs\wiki.md")
    targets = [
        Path(r"c:\Users\secha\Downloads\wiki.docx"),
        Path(r"c:\Users\secha\Downloads\OpsControl_Wiki.docx"),
        Path(r"c:\Users\secha\OpsControl\docs\wiki.docx"),
    ]
    convert_md_to_docx(src, targets)
